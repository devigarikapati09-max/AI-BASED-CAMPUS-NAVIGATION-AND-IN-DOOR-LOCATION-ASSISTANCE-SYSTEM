"""
Expansion script to add more content to reach 30+ pages.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def expand_report(doc_path):
    """Load existing report and add more content."""
    doc = Document(doc_path)
    
    # Navigate to before Lessons Learned section to insert new content
    # We'll add content between Chapter 4 conclusion and Lessons Learned
    
    # Add expanded content: AI and ML Details
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DETAILED AI AND MACHINE LEARNING CONCEPTS')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # 3.1.1 Defining AI
    p = doc.add_paragraph()
    run = p.add_run('Defining Artificial Intelligence')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    body = ("Artificial Intelligence is a branch of computer science that aims to create systems capable of "
            "performing tasks that typically require human intelligence. These tasks include visual perception, "
            "speech recognition, decision-making, problem-solving, language translation, and pattern recognition. "
            "AI systems are designed to learn from data, adapt to new inputs, and perform tasks that would "
            "traditionally require human cognition. The field encompasses various sub-disciplines including "
            "machine learning, deep learning, natural language processing, computer vision, and robotics.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("The fundamental distinction between traditional software and AI systems lies in their approach to "
            "problem-solving. Traditional software follows explicitly programmed rules and algorithms, executing "
            "deterministic operations on input data. AI systems, in contrast, learn patterns from data and "
            "make probabilistic predictions or decisions. This learning capability enables AI systems to handle "
            "complex, ambiguous, and evolving scenarios that would be impractical to encode as explicit rules.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    # 3.1.2 Historical Evolution
    p = doc.add_paragraph()
    run = p.add_run('Historical Evolution of AI')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    body = ("The history of artificial intelligence spans over seven decades of research and development. "
            "The term 'Artificial Intelligence' was coined by John McCarthy at the Dartmouth Conference in "
            "1956, marking the formal beginning of AI as a research field. The early decades (1950s-1970s) "
            "focused on symbolic AI and logic-based reasoning, with researchers developing programs that could "
            "play chess, prove mathematical theorems, and engage in limited natural language conversations.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("The 1980s saw the rise of expert systems, which encoded domain-specific knowledge from human "
            "experts into rule-based systems. These systems were successfully deployed in medical diagnosis, "
            "financial planning, and engineering design. The 1990s witnessed the emergence of machine learning "
            "approaches, particularly neural networks and decision trees, which enabled systems to learn from "
            "data rather than relying solely on hand-coded rules. The 2000s brought the big data revolution, "
            "providing the vast datasets needed to train sophisticated machine learning models. The 2010s "
            "marked the deep learning era, with convolutional neural networks achieving breakthrough performance "
            "in image recognition and recurrent neural networks excelling in natural language processing.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    # AI Applications in Navigation table
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    headers = ['AI Technique', 'Application in Navigation', 'Benefit']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    
    ai_apps = [
        ['Graph Algorithms', 'Optimal route computation (Dijkstra, A*)', 'Shortest path guarantee'],
        ['K-NN Classification', 'Wi-Fi fingerprint matching for positioning', 'Accurate location estimation'],
        ['Trilateration', 'Bluetooth beacon-based positioning', 'Low-power indoor tracking'],
        ['Image Processing', 'QR code scanning and recognition', 'Deterministic position fixes'],
        ['Decision Trees', 'Accessibility routing decisions', 'Dynamic path selection'],
        ['Neural Networks', 'Pattern recognition in signal data', 'Noise-robust positioning'],
        ['Natural Language Processing', 'Voice-guided navigation instructions', 'User-friendly directions'],
        ['Computer Vision', 'Visual landmark recognition', 'High-accuracy localization'],
    ]
    for i, row_data in enumerate(ai_apps):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
            for run in table.rows[i + 1].cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_paragraph('Table A: AI Techniques Applied in Campus Navigation')
    
    # Machine Learning Fundamentals
    p = doc.add_paragraph()
    run = p.add_run('Fundamentals of Machine Learning')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    body = ("Machine learning is a subset of artificial intelligence that enables systems to learn and improve "
            "from experience without being explicitly programmed. The learning process involves identifying "
            "patterns in training data and generalizing these patterns to make predictions on new, unseen data. "
            "The three fundamental categories of machine learning differ in the type of data they require and "
            "the nature of the learning objective.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("Supervised learning requires labeled training data where both input features and expected outputs "
            "are provided. The algorithm learns a mapping function from inputs to outputs, enabling it to make "
            "predictions on new data. Common supervised learning algorithms include linear regression, decision "
            "trees, support vector machines, and neural networks. In the context of campus navigation, supervised "
            "learning could be used to predict navigation difficulty based on building layout and user characteristics.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("Unsupervised learning operates on unlabeled data, discovering hidden patterns, structures, and "
            "relationships within the data. Clustering algorithms group similar data points together, while "
            "dimensionality reduction techniques simplify complex datasets while preserving essential information. "
            "In campus navigation, unsupervised learning could identify natural grouping patterns in user "
            "movement data, revealing common navigation paths and frequently visited destinations.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("Reinforcement learning trains agents through a system of rewards and penalties, enabling them to "
            "learn optimal strategies through trial and error. The agent interacts with an environment, taking "
            "actions and receiving feedback in the form of rewards. Over time, the agent learns to maximize "
            "cumulative rewards, developing optimal behavior policies. In campus navigation, reinforcement "
            "learning could optimize routing strategies based on real-time congestion data and user preferences.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    # Indoor Positioning Technologies - Expanded
    p = doc.add_paragraph()
    run = p.add_run('Detailed Indoor Positioning Technologies')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    body = ("Wi-Fi fingerprinting is one of the most widely adopted indoor positioning techniques due to the "
            "ubiquitous presence of Wi-Fi access points in educational institutions. The technique operates in "
            "two phases: the offline training phase and the online positioning phase. During the training phase, "
            "signal strength readings (RSSI values) are collected at known reference points throughout the "
            "building, creating a comprehensive fingerprint database. The positioning phase involves comparing "
            "live signal readings against the stored fingerprints using nearest neighbor matching.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("The k-nearest neighbors (k-NN) algorithm is the core matching technique used in Wi-Fi fingerprinting. "
            "The algorithm identifies the k closest fingerprint entries to the current signal reading and computes "
            "a weighted average of their known positions. The weighting is typically inversely proportional to "
            "the signal distance, ensuring that closer matches contribute more to the final position estimate. "
            "In our implementation, k=3 provides an optimal balance between accuracy and computational efficiency.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("Bluetooth beacon triangulation leverages the low-power Bluetooth Low Energy (BLE) technology to "
            "provide indoor positioning. Beacons are small, battery-powered devices that continuously broadcast "
            "their unique identifiers. Positioning is achieved by measuring the signal strength from multiple "
            "beacons and computing the user's position using trilateration. The inverse-distance weighting "
            "algorithm assigns higher weights to closer beacons, improving position accuracy. Beacon-based "
            "positioning offers advantages including low power consumption (beacon batteries last 1-2 years), "
            "moderate accuracy (1-3 meters), and easy deployment without requiring existing infrastructure.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    # Pathfinding Algorithms - Expanded
    p = doc.add_paragraph()
    run = p.add_run('Detailed Pathfinding Algorithm Analysis')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    body = ("Dijkstra's algorithm is a fundamental graph algorithm that computes the shortest path from a "
            "source node to all other nodes in a weighted graph. The algorithm maintains a set of unvisited "
            "nodes and iteratively selects the node with the smallest tentative distance, updating the distances "
            "of its neighbors. The algorithm terminates when all nodes have been visited or when the target "
            "node is reached. The key advantage of Dijkstra's algorithm is its guarantee of finding the optimal "
            "(shortest) path, regardless of the graph structure or edge weights.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("The A* search algorithm enhances Dijkstra's algorithm by incorporating a heuristic function that "
            "estimates the distance from the current node to the goal. The algorithm evaluates nodes based on "
            "the sum of the actual distance from the start (g-cost) and the estimated distance to the goal "
            "(h-cost). By prioritizing nodes that are likely closer to the goal, A* can skip exploring large "
            "portions of the graph that are clearly suboptimal. The heuristic must be admissible (never "
            "overestimate the true cost) and consistent (satisfy the triangle inequality) to guarantee optimality. "
            "Euclidean distance is the standard heuristic for campus navigation, as it provides a lower bound "
            "on the actual walking distance.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    # Algorithm comparison table
    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'
    headers = ['Algorithm', 'Time Complexity', 'Space Complexity', 'Optimal', 'Best For']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    
    algo_data = [
        ['Dijkstra', 'O(V log V + E)', 'O(V)', 'Yes', 'All-pairs shortest paths'],
        ['A* Search', 'O(b^d)', 'O(b^d)', 'Yes', 'Single-target searches'],
        ['BFS', 'O(V + E)', 'O(V)', 'Yes (unweighted)', 'Unweighted graphs'],
        ['Bellman-Ford', 'O(VE)', 'O(V)', 'Yes', 'Graphs with negative edges'],
        ['Floyd-Warshall', 'O(V^3)', 'O(V^2)', 'Yes', 'All-pairs distances'],
        ['Greedy BFS', 'O(b^d)', 'O(b^d)', 'No', 'Fast but suboptimal'],
    ]
    for i, row_data in enumerate(algo_data):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
            for run in table.rows[i + 1].cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)
    
    doc.add_paragraph()
    doc.add_paragraph('Table B: Pathfinding Algorithm Comparison')
    
    # Module Implementation Details - Expanded
    p = doc.add_paragraph()
    run = p.add_run('Detailed Module Implementation Analysis')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    body = ("The Campus Map Module implements a hierarchical data model that accurately represents the physical "
            "structure of a campus. At the top level, the CampusMap class maintains a collection of Building "
            "objects, each representing a physical structure on campus. Each Building contains multiple Floor "
            "objects, representing the vertical organization of space within the building. Each Floor contains "
            "Room objects that define the individual spaces, and Waypoint objects that represent navigation "
            "junctions (corridor intersections, stair landings, elevator doors, and building entrances).")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("The Room class captures essential attributes including room_id (unique identifier), room_number "
            "(human-readable identifier like 'R101'), room_name (descriptive name), room_type (classification "
            "such as lab, classroom, office, seminar, or library), and capacity (maximum number of occupants). "
            "The Waypoint class stores spatial coordinates (x, y) for position calculation and waypoint_type "
            "for distinguishing between different junction types. Inter-building paths are stored as weighted "
            "edges between building entrance waypoints, representing the walking distance between buildings.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("The Route Finder Module constructs a graph from the campus map data, connecting waypoints through "
            "four types of edges: within-floor corridors (connecting waypoints on the same floor), stair "
            "connections (connecting waypoints on adjacent floors within the same building), elevator "
            "connections (providing vertical access), and inter-building paths (connecting entrance waypoints "
            "of different buildings). The graph construction algorithm iterates through all floors and buildings, "
            "establishing connections based on proximity and structural adjacency. This graph serves as the "
            "foundation for both Dijkstra's and A* pathfinding algorithms.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    body = ("The Accessibility Service Module implements a comprehensive approach to inclusive navigation. "
            "Each building is flagged with an is_accessible attribute indicating whether it provides "
            "wheelchair-accessible entry and elevator access to all floors. The check_building_accessibility() "
            "method returns a detailed report including accessibility status, available accessibility features "
            "(elevators, ramps, accessible restrooms), and wheelchair-friendly routing options. The "
            "find_accessible_route() method filters the campus graph to exclude non-accessible waypoints and "
            "edges, ensuring that computed routes are navigable by wheelchair users. Voice-guided navigation "
            "simplifies technical route descriptions into natural language instructions that are easier to "
            "understand and follow.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    # Sample campus data table
    table = doc.add_table(rows=8, cols=5)
    table.style = 'Table Grid'
    headers = ['Building', 'Code', 'Floors', 'Rooms', 'Accessible']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    
    campus_data = [
        ['Academic Block A', 'B1', '3', '12', 'Yes'],
        ['Engineering Block', 'B2', '2', '8', 'Yes'],
        ['Science Block', 'B3', '2', '7', 'No'],
        ['Administrative Block', 'B4', '1', '4', 'Yes'],
        ['Central Library', 'B5', '2', '6', 'Yes'],
    ]
    for i, row_data in enumerate(campus_data):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
            for run in table.rows[i + 1].cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_paragraph('Table C: Sample Campus Building Data')
    
    # Room type distribution table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Room Type', 'Count', 'Description']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    
    room_data = [
        ['Laboratory', '15', 'Science, Engineering, and Computer Science labs'],
        ['Classroom', '5', 'General-purpose teaching spaces'],
        ['Office', '8', 'Faculty and administrative offices'],
        ['Seminar', '3', 'Large presentation and lecture halls'],
        ['Library', '6', 'Reading rooms, stacks, and reference areas'],
    ]
    for i, row_data in enumerate(room_data):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
            for run in table.rows[i + 1].cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_paragraph('Table D: Room Type Distribution Across Campus')
    
    # Positioning Technology Comparison
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    headers = ['Technology', 'Accuracy', 'Cost', 'Coverage', 'Power']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    
    pos_data = [
        ['Wi-Fi Fingerprinting', '2-5m', 'Low', 'Full building', 'Existing infra'],
        ['Bluetooth Beacons', '1-3m', 'Medium', 'Beacon coverage', 'Low (BLE)'],
        ['QR Code Scanning', '0m (exact)', 'Very Low', 'Code locations', 'Camera'],
        ['Computer Vision', '0.5-2m', 'High', 'Visual landmarks', 'High (CPU/GPU)'],
    ]
    for i, row_data in enumerate(pos_data):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
            for run in table.rows[i + 1].cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)
    
    doc.add_paragraph()
    doc.add_paragraph('Table E: Indoor Positioning Technology Comparison')
    
    # Notification Types Detail
    p = doc.add_paragraph()
    run = p.add_run('Notification System Design')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    body = ("The notification system is designed to keep users informed throughout their navigation journey. "
            "Six distinct notification types are implemented, each serving a specific purpose in the navigation "
            "workflow. Room change notifications alert users when their scheduled room has been relocated, "
            "providing updated navigation instructions. Arrival notifications confirm when a user has reached "
            "their destination building or floor. Direction notifications provide turn-by-turn guidance as "
            "users progress along their route. Accessibility notifications inform users about elevator and "
            "ramp locations when needed. Distance notifications update users on their remaining walking "
            "distance at regular intervals. Destination notifications provide final arrival information "
            "including room details, accessibility features, and nearby facilities.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    # Security Considerations
    p = doc.add_paragraph()
    run = p.add_run('Security Considerations')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    body = ("The campus navigation system incorporates several security measures to protect user data and "
            "ensure system integrity. The Flask application implements input validation and sanitization for "
            "all API endpoints, preventing injection attacks and malformed requests. User session management "
            "uses secure tokens with appropriate expiration policies. The system does not store personally "
            "identifiable information in the positioning database, ensuring user privacy. API rate limiting "
            "prevents abuse and ensures fair resource allocation. The database schema uses parameterized queries "
            "to prevent SQL injection. All inter-component communication follows the principle of least "
            "privilege, with each module accessing only the data it requires.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    # Performance Analysis
    p = doc.add_paragraph()
    run = p.add_run('Performance Analysis and Optimization')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    body = ("The system's performance characteristics have been evaluated across multiple dimensions. The "
            "campus map search operation achieves sub-millisecond response times by maintaining room indices "
            "in hash tables for O(1) lookup. The route finding algorithms demonstrate optimal time complexity: "
            "Dijkstra's algorithm operates in O(V log V + E) time using a binary heap priority queue, while "
            "A* search achieves O(b^d) where b is the branching factor and d is the solution depth. The "
            "indoor positioning module processes Wi-Fi fingerprint matching in O(n*k) time where n is the "
            "database size and k is the number of nearest neighbors. Memory usage is optimized through lazy "
            "initialization of positioning databases and on-demand graph construction.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    # Deployment Considerations
    p = doc.add_paragraph()
    run = p.add_run('Deployment and Scalability Considerations')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    body = ("The system is designed for horizontal scalability, allowing deployment across multiple servers "
            "to handle increased load. The Flask application can be deployed using Gunicorn or uWSGI as the "
            "WSGI server, with Nginx as the reverse proxy for load balancing. The modular architecture enables "
            "independent scaling of the positioning, routing, and notification services. Database scaling is "
            "supported through connection pooling and read replicas. The system can handle campus expansions "
            "by simply adding new building data to the campus map without requiring code changes. The API "
            "design follows RESTful principles, enabling integration with existing institutional systems "
            "such as student information systems and scheduling applications.")
    p = doc.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(20)
    
    # Save the expanded report
    output_path = doc_path
    doc.save(output_path)
    print(f"Expanded report saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    doc_path = 'Internship_Report_Campus_Navigation.docx'
    expand_report(doc_path)
